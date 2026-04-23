#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SQL Server 数据库巡检报告生成工具
适配 Windows 服务器版本
主要改动：
  - 使用 psutil 获取主机资源信息，替代 Linux 特有命令
  - 保留原有所有功能，数据格式不变
"""
from __future__ import unicode_literals
import argparse
import pyodbc
import configparser
import docx
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import logging
import logging.handlers
import time
import datetime as dt
from datetime import datetime, timedelta
import json
import os
import sys
import decimal
import shutil
import platform
import subprocess
import hashlib
import base64
import traceback

# 尝试导入psutil（Windows下用于获取主机资源）
try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False
    # 后面会打印警告

# 尝试导入openpyxl（用于Excel操作）
try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# 导入许可证管理（保持原有逻辑）
from license_manager import LicenseValidator

# 路径工具函数：适配开发和打包环境
def get_resource_path(relative_path):
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)

# 炫酷输出类（保持原有）
class CoolPrinter:
    def __init__(self):
        self.terminal_width = shutil.get_terminal_size().columns
        self.colors = {
            'green': '\033[92m',
            'yellow': '\033[93m',
            'blue': '\033[94m',
            'magenta': '\033[95m',
            'cyan': '\033[96m',
            'red': '\033[91m',
            'bold': '\033[1m',
            'end': '\033[0m'
        }
    
    def print_step(self, step_num, title, status=None):
        if status == "start":
            print(f"\n{self.colors['bold']}{self.colors['blue']}步骤{step_num}: {title}{self.colors['end']}")
            print("=" * self.terminal_width)
        elif status == "success":
            print(f"{self.colors['green']}✓ {title}{self.colors['end']}")
        elif status == "error":
            print(f"{self.colors['red']}✗ {title}{self.colors['end']}")
        elif status == "warning":
            print(f"{self.colors['yellow']}⚠ {title}{self.colors['end']}")
        else:
            print(f"{self.colors['cyan']}➤ {title}{self.colors['end']}")
    
    def print_progress(self, current, total, prefix="", suffix="", length=30):
        percent = current / total
        filled = int(length * percent)
        bar = '█' * filled + '░' * (length - filled)
        progress_str = f"{prefix} |{bar}| {current}/{total} {suffix}"
        print(f"\r{progress_str}", end='', flush=True)
        if current == total:
            print()
    
    def print_info(self, message):
        print(f"{self.colors['cyan']}ℹ {message}{self.colors['end']}")
    
    def print_success(self, message):
        print(f"{self.colors['green']}✓ {message}{self.colors['end']}")
    
    def print_warning(self, message):
        print(f"{self.colors['yellow']}⚠ {message}{self.colors['end']}")
    
    def print_error(self, message):
        print(f"{self.colors['red']}✗ {message}{self.colors['end']}")

printer = CoolPrinter()

# 日志配置
def getlogger():
    logger = logging.getLogger()
    if not logger.handlers:
        logger.setLevel(logging.DEBUG)
        ch = logging.FileHandler(r'autoDoc.log')
        ch.setLevel(logging.DEBUG)
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        ch.setFormatter(formatter)
        logger.addHandler(ch)
    return logger

logger = getlogger()

# 参数解析类
class passArgu(object):
    def get_argus(self):
        all_info = argparse.ArgumentParser(
            description="SQL Server 数据库巡检报告生成工具")
        all_info.add_argument('-F', '--conf', required=False, default='conf/sqlserver.ini', help='SQL Server config file (INI format).')
        all_info.add_argument('-C', '--sqltemplates', required=False, default='templates/sqltemplates_sqlserver.ini', help='SQL templates file.')
        all_info.add_argument('-L', '--label', required=False, help='Label used when health check single database.')
        all_info.add_argument('-E', '--excel-connections', required=False, help='Excel file containing database connection information for batch inspection.')
        all_info.add_argument('-G', '--generate-excel-template', required=False, metavar='TEMPLATE_FILE', help='Generate an Excel template for connection information and exit.')
        all_info.add_argument('-S', '--summary-report', required=False, action='store_true', help='Generate a summary Excel report after batch inspection.')
        all_para = all_info.parse_args()
        return all_para

# ==================== 默认SQL模板内容（增强兼容性，避免SQL_VARIANT错误） ====================
DEFAULT_SQL_TEMPLATES = {
    "variables": {
        "version": "SELECT @@VERSION AS version;",
        "product_version": "SELECT CAST(SERVERPROPERTY('ProductVersion') AS NVARCHAR(128)) AS product_version;",
        "edition": "SELECT CAST(SERVERPROPERTY('Edition') AS NVARCHAR(128)) AS edition;",
        "product_level": "SELECT CAST(SERVERPROPERTY('ProductLevel') AS NVARCHAR(128)) AS product_level;",
        "max_connections": "SELECT @@MAX_CONNECTIONS AS max_connections;",
        "active_users": "SELECT session_id, login_name, host_name, program_name, status FROM sys.dm_exec_sessions WHERE is_user_process = 1;",
        "db_size": """
            SELECT 
                d.name AS database_name,
                SUM(mf.size) * 8 / 1024.0 AS size_mb,
                d.state_desc,
                d.recovery_model_desc
            FROM sys.databases d
            JOIN sys.master_files mf ON d.database_id = mf.database_id
            GROUP BY d.name, d.state_desc, d.recovery_model_desc;
        """,
        "disk_space": "EXEC xp_fixeddrives;",
        "config_parameters": """
            SELECT name, 
                   CAST(value AS NVARCHAR(100)) AS value, 
                   CAST(value_in_use AS NVARCHAR(100)) AS value_in_use
            FROM sys.configurations
            WHERE name IN (
                'max server memory (MB)', 
                'min server memory (MB)', 
                'max degree of parallelism',
                'cost threshold for parallelism',
                'backup compression default'
            );
        """,
        "database_status": "SELECT name, state_desc, recovery_model_desc, compatibility_level FROM sys.databases;",
        "wait_stats": """
            SELECT TOP 10 wait_type, wait_time_ms, waiting_tasks_count
            FROM sys.dm_os_wait_stats
            WHERE wait_time_ms > 0
            ORDER BY wait_time_ms DESC;
        """,
        "cpu_usage": "SELECT @@CPU_BUSY AS cpu_busy_ms, @@IDLE AS idle_ms, @@TIMETICKS AS time_ticks;",
        "blocking": """
            SELECT session_id, blocking_session_id, wait_duration_ms, wait_type
            FROM sys.dm_os_waiting_tasks
            WHERE blocking_session_id <> 0;
        """,
        "all_databases": "SELECT name, state_desc, recovery_model_desc, compatibility_level FROM sys.databases;",
        "log_space": """
            -- 使用 DBCC SQLPERF 获取日志空间信息（所有版本兼容）
            CREATE TABLE #logspace (
                DatabaseName sysname,
                LogSizeMB float,
                LogSpaceUsedPercent float,
                Status int
            );
            INSERT #logspace EXEC ('DBCC SQLPERF(LOGSPACE)');
            SELECT DatabaseName, LogSizeMB, LogSpaceUsedPercent, Status FROM #logspace;
            DROP TABLE #logspace;
        """,
        "memory_usage": """
            -- 使用 sys.dm_os_process_memory (SQL Server 2008+)
            SELECT 
                'Total' AS type,
                CAST(physical_memory_in_use_kb/1024.0 AS DECIMAL(10,2)) AS memory_used_mb,
                CAST(available_commit_limit_kb/1024.0 AS DECIMAL(10,2)) AS available_mb
            FROM sys.dm_os_process_memory;
        """,
        "linked_servers": "EXEC sp_linkedservers;",
        "file_groups": """
            SELECT 
                data_space_id,
                name,
                type_desc,
                is_default
            FROM sys.filegroups;
        """,
        "backup_info": """
            -- 获取每个数据库的最近备份信息
            SELECT 
                database_name,
                MAX(backup_finish_date) AS last_backup_time,
                CASE MAX(type) WHEN 'D' THEN '完整' WHEN 'I' THEN '差异' WHEN 'L' THEN '日志' ELSE '其他' END AS backup_type,
                CAST(MAX(backup_size)/1024.0/1024.0 AS DECIMAL(10,2)) AS backup_size_mb
            FROM msdb.dbo.backupset
            GROUP BY database_name
            ORDER BY database_name;
        """,
        "locks": """
            -- 当前锁数量和最长等待时间
            SELECT 
                (SELECT COUNT(*) FROM sys.dm_tran_locks) AS lock_count,
                ISNULL(MAX(wt.wait_duration_ms), 0) AS max_wait_time
            FROM sys.dm_os_waiting_tasks wt
            JOIN sys.dm_tran_locks tl ON wt.resource_address = tl.lock_owner_address
            WHERE wt.wait_type LIKE 'LCK%';
        """,
        "deadlocks": """
            -- 过去24小时死锁数量（从扩展事件文件读取，可能因配置失败，失败时返回0）
            SELECT COUNT(*) AS deadlock_count
            FROM sys.fn_xe_file_target_read_file('DeadlockTracking*.xel', NULL, NULL, NULL)
            WHERE object_name = 'xml_deadlock_report';
        """,
    },
    "report": {
        "template": "templates/report_template.docx",
        "output": "reports/{label}_report_{timestamp}.docx"
    }
}

def init_sql_templates(template_file):
    """初始化SQL模板文件，始终覆盖创建最新版本。"""
    printer.print_info(f"正在初始化SQL模板: {template_file}")
    os.makedirs(os.path.dirname(template_file), exist_ok=True)
    config = configparser.RawConfigParser()
    for section, items in DEFAULT_SQL_TEMPLATES.items():
        config.add_section(section)
        for key, value in items.items():
            config.set(section, key, value.strip())
    with open(template_file, 'w', encoding='utf-8') as f:
        config.write(f)
    printer.print_success(f"SQL模板已更新: {template_file}")

# ==================== Excel操作相关函数 ====================
def generate_excel_template(template_file):
    if not OPENPYXL_AVAILABLE:
        printer.print_error("openpyxl 模块未安装，无法生成Excel模板。请运行: pip install openpyxl")
        return False
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "数据库连接"
        headers = ['Label', 'Server', 'Port', 'Database', 'User', 'Password', 'Driver']
        ws.append(headers)
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
            cell.alignment = Alignment(horizontal='center')
        example = [
            ['TEST_DB', '192.168.1.100', '1433', 'master', 'sa', 'your_password', '{ODBC Driver 17 for SQL Server}'],
            ['PROD_DB', 'sqlserver.company.com', '1433', 'AdventureWorks', 'inspector', 'inspector_pwd', '{SQL Server}'],
        ]
        for row in example:
            ws.append(row)
        for col in ws.columns:
            max_len = 0
            col_letter = col[0].column_letter
            for cell in col:
                try:
                    if cell.value:
                        max_len = max(max_len, len(str(cell.value)))
                except:
                    pass
            adjusted_width = min(max_len + 2, 50)
            ws.column_dimensions[col_letter].width = adjusted_width
        wb.save(template_file)
        printer.print_success(f"Excel模板已生成: {template_file}")
        return True
    except Exception as e:
        printer.print_error(f"生成Excel模板失败: {e}")
        return False

def read_connections_from_excel(excel_file):
    if not OPENPYXL_AVAILABLE:
        printer.print_error("openpyxl 模块未安装，无法读取Excel文件。")
        return []
    try:
        wb = load_workbook(excel_file)
        ws = wb.active
        headers = [cell.value for cell in ws[1]]
        required = {'Label', 'Server', 'Port', 'Database', 'User', 'Password', 'Driver'}
        if not required.issubset(set(headers)):
            missing = required - set(headers)
            printer.print_error(f"Excel文件缺少必要的列: {missing}")
            return []
        connections = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            row_dict = dict(zip(headers, row))
            if not any(row_dict.values()):
                continue
            conn = {k: str(v).strip() if v is not None else '' for k, v in row_dict.items()}
            connections.append(conn)
        printer.print_success(f"从Excel读取到 {len(connections)} 个数据库连接")
        return connections
    except Exception as e:
        printer.print_error(f"读取Excel文件失败: {e}")
        return []

def generate_summary_excel(summary_data, output_file):
    if not OPENPYXL_AVAILABLE:
        printer.print_error("openpyxl 未安装，无法生成汇总报告")
        return False
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "巡检汇总"
        headers = ['标签', '服务器', '数据库', '版本', '总大小(MB)', '活动会话数', 
                   '最大连接数', '检查时间', '报告文件']
        ws.append(headers)
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
            cell.alignment = Alignment(horizontal='center')
        for item in summary_data:
            row = [
                item.get('label', ''),
                item.get('server', ''),
                item.get('database', ''),
                item.get('version', ''),
                item.get('total_size_mb', ''),
                item.get('active_sessions', ''),
                item.get('max_connections', ''),
                item.get('check_time', ''),
                item.get('report_file', '')
            ]
            ws.append(row)
        for col in ws.columns:
            max_len = 0
            col_letter = col[0].column_letter
            for cell in col:
                if cell.value:
                    max_len = max(max_len, len(str(cell.value)))
            ws.column_dimensions[col_letter].width = min(max_len + 2, 50)
        wb.save(output_file)
        printer.print_success(f"汇总报告已生成: {output_file}")
        return True
    except Exception as e:
        printer.print_error(f"生成汇总报告失败: {e}")
        return False

# ==================== 数据获取类 ====================
class connInfo(object):
    def __init__(self):
        self.printer = printer
        infos = passArgu().get_argus()
        self.jdbcinfo = get_resource_path(str(infos.conf))
        self.label = infos.label

    def jdbcInfo(self):
        self.printer.print_step(1, "数据库连接配置", "start")
        self.printer.print_info("检查ODBC驱动安装状态...")
        try:
            drivers = pyodbc.drivers()
            sql_server_drivers = [d for d in drivers if 'SQL Server' in d]
            if sql_server_drivers:
                for driver in sql_server_drivers:
                    self.printer.print_success(f"找到驱动: {driver}")
            else:
                self.printer.print_error("未找到SQL Server ODBC驱动")
                sys.exit(1)
        except Exception as e:
            self.printer.print_error(f"检查ODBC驱动失败: {e}")
            sys.exit(1)
        
        conf = self.jdbcinfo
        label = self.label
        cfg = configparser.RawConfigParser()
        if not os.path.exists(conf):
            self.printer.print_error(f"{conf} 配置文件不存在")
            sys.exit(2)
        try:
            cfg.read(conf)
        except Exception as e:
            self.printer.print_error(f"读取 {conf} 配置文件时出错: {e}")
            sys.exit(2)

        dbinfo = {}
        if label:
            if cfg.has_section(label):
                try:
                    name = cfg.get(label, "name")
                    server = cfg.get(label, "server")
                    port = cfg.get(label, "port")
                    user = cfg.get(label, "user")
                    password = cfg.get(label, "password")
                    driver = cfg.get(label, "driver")
                    dbinfo[label] = {"name": name, "server": server, "port": port, "user": user, "password": password, "driver": driver}
                except configparser.NoOptionError as e:
                    self.printer.print_error(f"配置文件 {label} 部分缺少选项: {e}")
                    sys.exit(2)
            else:
                self.printer.print_error(f"配置文件中不存在 {label} 部分")
                sys.exit(2)
        else:
            for labels in cfg.sections():
                try:
                    name = cfg.get(labels, "name")
                    server = cfg.get(labels, "server")
                    port = cfg.get(labels, "port")
                    user = cfg.get(labels, "user")
                    password = cfg.get(labels, "password")
                    driver = cfg.get(labels, "driver")
                    dbinfo[labels] = {"name": name, "server": server, "port": port, "user": user, "password": password, "driver": driver}
                except configparser.NoOptionError as e:
                    self.printer.print_error(f"配置文件 {labels} 部分缺少选项: {e}")
        return dbinfo

class getData(object):
    def __init__(self, server, port, user, password, driver, label):
        self.printer = printer
        self.label = label
        self.server = server
        self.port = port
        self.user = user
        self.password = password

        available_drivers = pyodbc.drivers()
        clean_driver = driver.strip('{}')
        if clean_driver not in available_drivers:
            matched = [d for d in available_drivers if 'SQL Server' in d]
            if matched:
                self.printer.print_warning(f"指定的驱动 '{driver}' 不存在，将使用 '{matched[0]}'")
                driver = matched[0]
            else:
                self.printer.print_error("未找到任何 SQL Server ODBC 驱动，请检查安装。")
                raise Exception("No SQL Server ODBC driver found")
        else:
            driver = clean_driver

        self.conn_str = f'DRIVER={{{driver}}};SERVER={server},{port};UID={user};PWD={password};DATABASE=master'
        try:
            self.conn = pyodbc.connect(self.conn_str)
            self.printer.print_success(f"已连接到服务器: {server}")
        except pyodbc.Error as e:
            self.printer.print_error(f"数据库连接失败: {e}")
            try:
                self.conn_str = f'DRIVER={driver};SERVER={server},{port};UID={user};PWD={password};DATABASE=master'
                self.conn = pyodbc.connect(self.conn_str)
                self.printer.print_success(f"使用不带花括号的驱动名称连接成功: {server}")
            except Exception as e2:
                self.printer.print_error(f"再次连接失败: {e2}")
                raise

        self.context = {}

    def checkdb(self, sqlfile=''):
        if not os.path.exists(sqlfile):
            self.printer.print_error(f"{sqlfile} SQL模板文件不存在")
            raise FileNotFoundError(f"{sqlfile} SQL模板文件不存在")
        cfg = configparser.RawConfigParser()
        try:
            cfg.read(sqlfile, encoding='utf-8')
        except Exception as e:
            self.printer.print_error(f"解析 {sqlfile} 配置文件时出错: {e}")
            raise

        self.context = {
            "checktime": [], "version": [], "db_size": [], "max_connections": [],
            "active_users": [], "os_parameters": [], "start_parameters": [],
            "start_time": [], "server_name": [], "instance_name": [], "disk_space": [],
            "log_file_info": [], "table_disk_space": [], "io_work": [], "cpu_activity": [],
            "locks": [], "deadlocks": [], "active_users_processes": [], "user_roles": [],
            "linked_servers": [], "file_groups_files": [], "memory_usage": [],
            "log_space_info": [], "total_server_memory": [], "sjcp": [],
            "table_space_usage": [], "all_databases": [], "log": [],
            "product_version": [], "edition": [], "product_level": [], "config_parameters": [],
            "database_status": [], "wait_stats": [], "cpu_usage": [], "blocking": [],
            "log_space": [], "file_groups": [], "backup_info": [], "locks": [], "deadlocks": []
        }

        all_queries = list(cfg.items("variables"))
        total_queries = len(all_queries)
        self.printer.print_info("开始执行数据库巡检...")
        try:
            cursor = self.conn.cursor()
            for idx, (name, stmt) in enumerate(all_queries, 1):
                try:
                    cursor.execute(stmt)
                    if cursor.description:
                        columns = [column[0] for column in cursor.description]
                        rows = cursor.fetchall()
                        r = []
                        for row in rows:
                            row_dict = {}
                            for col_idx, col_name in enumerate(columns):
                                value = row[col_idx]
                                # 处理特殊类型
                                if isinstance(value, decimal.Decimal):
                                    value = float(value)
                                elif isinstance(value, dt.datetime):
                                    value = value.strftime('%Y-%m-%d %H:%M:%S')
                                elif isinstance(value, dt.date):
                                    value = value.strftime('%Y-%m-%d')
                                row_dict[col_name] = value
                            r.append(row_dict)
                        self.context[name] = r
                    else:
                        self.context[name] = []
                    self.printer.print_progress(idx, total_queries, prefix=f"收集数据 {idx}/{total_queries}", suffix=name)
                except Exception as e:
                    self.printer.print_error(f"执行 {name} 查询时出错: {e}")
                    self.context[name] = []
        except Exception as e:
            self.printer.print_error(f"执行查询时出错: {e}")
        finally:
            try:
                cursor.close()
            except:
                pass
        return self.context

# ==================== 主机资源获取函数（Windows 适配，使用 psutil） ====================
def get_host_disk_usage():
    """
    获取主机磁盘使用情况（Windows 版）
    返回格式: {"data": [ [device, mountpoint, fstype, total_gb, used_gb, free_gb, usage_percent], ... ]}
    """
    if not PSUTIL_AVAILABLE:
        printer.print_warning("psutil 未安装，无法获取磁盘使用率")
        return {"data": []}
    try:
        disk_data = []
        partitions = psutil.disk_partitions()
        for part in partitions:
            try:
                usage = psutil.disk_usage(part.mountpoint)
                total_gb = round(usage.total / (1024**3), 2)
                used_gb = round(usage.used / (1024**3), 2)
                free_gb = round(usage.free / (1024**3), 2)
                usage_percent = usage.percent
                # device 在 Windows 下通常是盘符如 "C:"，mountpoint 也是 "C:\\"
                disk_data.append([
                    part.device,
                    part.mountpoint,
                    part.fstype,
                    total_gb,
                    used_gb,
                    free_gb,
                    usage_percent
                ])
            except PermissionError:
                # 某些挂载点可能无法访问（如 CD-ROM）
                continue
        return {"data": disk_data}
    except Exception as e:
        printer.print_warning(f"获取磁盘使用率失败: {str(e)}")
        return {"data": []}

def get_host_memory_usage():
    """
    获取主机内存使用情况（Windows 版）
    返回格式: {"data": [ ["物理内存", total_gb, used_gb, free_gb, usage_percent] ]}
    """
    if not PSUTIL_AVAILABLE:
        printer.print_warning("psutil 未安装，无法获取内存使用率")
        return {"data": []}
    try:
        mem = psutil.virtual_memory()
        total_gb = round(mem.total / (1024**3), 2)
        # used = total - available (与 Linux 原方法一致)
        used_gb = round((mem.total - mem.available) / (1024**3), 2)
        free_gb = round(mem.available / (1024**3), 2)
        usage_percent = mem.percent
        memory_data = [["物理内存", total_gb, used_gb, free_gb, usage_percent]]
        return {"data": memory_data}
    except Exception as e:
        printer.print_warning(f"获取内存使用率失败: {str(e)}")
        return {"data": []}

def get_host_cpu_usage():
    """
    获取主机 CPU 使用情况（Windows 版）
    返回格式: {"data": [ [1, "xx.x%"] ]}   # 第一个元素固定为1，第二个为使用率字符串
    """
    if not PSUTIL_AVAILABLE:
        printer.print_warning("psutil 未安装，无法获取CPU使用率")
        return {"data": []}
    try:
        # 采样0.1秒，避免长时间阻塞
        cpu_percent = psutil.cpu_percent(interval=0.1)
        usage_str = f"{cpu_percent:.1f}%"
        return {"data": [[1, usage_str]]}
    except Exception as e:
        printer.print_warning(f"获取CPU使用率失败: {str(e)}")
        return {"data": []}

# ==================== 巡检结果分析函数（添加类型安全，修改告警逻辑）====================
def analyze_inspection_results(result):
    stats = {
        'tablespace_count': 0,
        'urgent_expansion': 0,
        'need_attention': 0,
        'alarm_count': 0,
        'critical_alarm': 0,
        'memory_usage_rate': 0.0,
        'cpu_usage_rate': 0.0,
        'disk_usage_rate': 0.0,
        'host_memory_usage': 0.0,
        'host_cpu_usage': 0.0,
        'health_status': "正常",
        'health_desc': "数据库运行状态良好，无明显异常"
    }
    
    # 安全获取并处理每个字段
    def safe_list_get(data):
        if isinstance(data, list):
            return data
        return []
    
    db_size_data = safe_list_get(result.get('db_size'))
    stats['tablespace_count'] = len(db_size_data)

    wait_stats = safe_list_get(result.get('wait_stats'))
    if len(wait_stats) > 0:
        stats['alarm_count'] += 1  # 等待统计存在，计入普通告警，不计严重
        # 原代码同时增加critical_alarm，现移除

    blocking = safe_list_get(result.get('blocking'))
    if len(blocking) > 0:
        stats['alarm_count'] += len(blocking)  # 每个阻塞进程计入普通告警
        # 原代码同时增加critical_alarm，现移除

    max_conn = safe_list_get(result.get('max_connections'))
    if max_conn and 'max_connections' in max_conn[0]:
        stats['max_connections'] = max_conn[0]['max_connections']

    active = safe_list_get(result.get('active_users'))
    stats['active_sessions'] = len(active)

    # 主机资源
    host_disk = result.get('host_disk_usage', {})
    if isinstance(host_disk, dict):
        disk_data = safe_list_get(host_disk.get('data'))
    else:
        disk_data = []
    if disk_data:
        # 取所有分区的最大使用率，超过95%才视为严重
        max_disk = max([d[6] for d in disk_data if len(d) > 6 and isinstance(d[6], (int, float))], default=0)
        stats['disk_usage_rate'] = max_disk
        if max_disk >= 95:  # 阈值从90改为95
            stats['alarm_count'] += 1
            stats['critical_alarm'] += 1

    host_mem = result.get('host_memory_usage', {})
    if isinstance(host_mem, dict):
        mem_data = safe_list_get(host_mem.get('data'))
    else:
        mem_data = []
    if mem_data and len(mem_data[0]) >= 5:
        stats['host_memory_usage'] = mem_data[0][4]
        if mem_data[0][4] >= 95:  # 阈值从90改为95
            stats['alarm_count'] += 1
            stats['critical_alarm'] += 1

    host_cpu = result.get('host_cpu_usage', {})
    if isinstance(host_cpu, dict):
        cpu_data = safe_list_get(host_cpu.get('data'))
    else:
        cpu_data = []
    if cpu_data and len(cpu_data[0]) >= 2:
        cpu_str = cpu_data[0][1]
        try:
            cpu_val = float(cpu_str.replace('%', ''))
            stats['host_cpu_usage'] = cpu_val
            if cpu_val >= 95:  # 阈值从90改为95
                stats['alarm_count'] += 1
                stats['critical_alarm'] += 1
        except:
            pass

    # 锁和死锁
    locks = safe_list_get(result.get('locks'))
    if locks and 'lock_count' in locks[0]:
        if locks[0].get('lock_count', 0) > 50:
            stats['alarm_count'] += 1  # 锁数量过多，计入普通告警，不计严重

    deadlocks = safe_list_get(result.get('deadlocks'))
    if deadlocks and 'deadlock_count' in deadlocks[0]:
        if deadlocks[0].get('deadlock_count', 0) > 0:
            stats['alarm_count'] += 1  # 死锁存在，计入普通告警，不计严重
            # 原代码同时增加critical_alarm，现移除

    if stats['critical_alarm'] > 0:
        stats['health_status'] = "严重"
        stats['health_desc'] = f"发现{stats['critical_alarm']}个严重告警，请立即处理"
    elif stats['alarm_count'] > 0:
        stats['health_status'] = "警告"
        stats['health_desc'] = f"发现{stats['alarm_count']}个告警，建议及时排查"
    else:
        stats['health_status'] = "正常"
        stats['health_desc'] = "数据库运行状态良好"
    return stats

# ==================== 创建Word报告（动态构建，全面类型安全）====================
def set_chinese_font(paragraph, font_size=Pt(10)):
    try:
        if paragraph.runs:
            run = paragraph.runs[0]
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = font_size
    except:
        pass

def set_table_style(table):
    try:
        table.style = 'Table Grid'
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                if cell.paragraphs:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold = True
        for row in table.rows[1:]:
            for cell in row.cells:
                if cell.paragraphs:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        pass

def safe_list(data):
    """将数据安全地转换为列表，如果不是列表则返回空列表"""
    if isinstance(data, list):
        return data
    return []

def safe_dict(data):
    """安全地获取字典，如果不是字典则返回空字典"""
    if isinstance(data, dict):
        return data
    return {}

def safe_get(data, key, default=None):
    """安全地从字典中获取值，如果 data 不是字典则返回 default"""
    if isinstance(data, dict):
        return data.get(key, default)
    return default

def create_sqlserver_report(db_info, result, inspector="", system=""):
    doc = docx.Document()
    # 标题字体改为小二（18pt），原为28pt
    title = doc.add_heading('', level=0)
    title_run = title.add_run("SQL Server 数据库巡检报告")
    title_run.font.size = Pt(18)  # 修改为18pt
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_chinese_font(title)

    # 报告基本信息
    doc.add_heading('一、报告基本信息', level=1)
    info_table = doc.add_table(rows=7, cols=2)
    info_table.style = 'Table Grid'
    rows = [
        ('数据库标签', db_info.get('label', '')),
        ('服务器', f"{db_info.get('server', '')}:{db_info.get('port', '')}"),
        ('数据库', db_info.get('name', '')),
        ('检查时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ('连接用户', db_info.get('user', '')),
        ('巡检人', inspector),
        ('巡检系统', system)
    ]
    for i, (label, value) in enumerate(rows):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = str(value)  # 确保字符串
    set_table_style(info_table)

    # 二、数据库版本信息（表格形式）
    doc.add_heading('二、数据库版本信息', level=1)
    version_table = doc.add_table(rows=4, cols=2)
    version_table.style = 'Table Grid'
    version_list = safe_list(result.get('version'))
    product_version_list = safe_list(result.get('product_version'))
    edition_list = safe_list(result.get('edition'))
    product_level_list = safe_list(result.get('product_level'))
    version_items = [
        ('SQL Server 版本', version_list[0].get('version', 'N/A')[:200] if version_list else 'N/A'),
        ('产品版本', product_version_list[0].get('product_version', 'N/A') if product_version_list else 'N/A'),
        ('版本', edition_list[0].get('edition', 'N/A') if edition_list else 'N/A'),
        ('Service Pack', product_level_list[0].get('product_level', 'N/A') if product_level_list else 'N/A')
    ]
    for i, (key, val) in enumerate(version_items):
        version_table.cell(i, 0).text = key
        version_table.cell(i, 1).text = str(val)  # 确保字符串
    set_table_style(version_table)

    # 三、连接数信息（表格）
    doc.add_heading('三、连接数信息', level=1)
    conn_table = doc.add_table(rows=2, cols=2)
    conn_table.style = 'Table Grid'
    max_conn_list = safe_list(result.get('max_connections'))
    max_val = max_conn_list[0].get('max_connections', 'N/A') if max_conn_list else 'N/A'
    active_users_list = safe_list(result.get('active_users'))
    active_sessions = len(active_users_list)
    conn_items = [
        ('最大连接数', max_val),
        ('当前活动会话数', str(active_sessions))  # active_sessions 已转字符串，但为统一也转为字符串
    ]
    for i, (key, val) in enumerate(conn_items):
        conn_table.cell(i, 0).text = key
        conn_table.cell(i, 1).text = str(val)  # 关键修复：强制转为字符串
    set_table_style(conn_table)

    # 四、数据库大小信息
    doc.add_heading('四、数据库大小信息', level=1)
    db_size = safe_list(result.get('db_size'))
    if db_size:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = '数据库名称'
        hdr[1].text = '大小(MB)'
        hdr[2].text = '状态'
        hdr[3].text = '恢复模式'
        for item in db_size:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('database_name', ''))
            row_cells[1].text = str(round(item.get('size_mb', 0), 2))  # 转为字符串
            row_cells[2].text = str(item.get('state_desc', ''))
            row_cells[3].text = str(item.get('recovery_model_desc', ''))
        set_table_style(table)
    else:
        doc.add_paragraph("无数据库大小信息")

    # 五、关键配置参数
    doc.add_heading('五、关键配置参数', level=1)
    config_params = safe_list(result.get('config_parameters'))
    if config_params:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = '参数名'
        hdr[1].text = '配置值'
        hdr[2].text = '运行值'
        for item in config_params:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('name', ''))
            row_cells[1].text = str(item.get('value', ''))
            row_cells[2].text = str(item.get('value_in_use', ''))
        set_table_style(table)
    else:
        doc.add_paragraph("无配置参数信息")

    # 六、活动用户会话
    doc.add_heading('六、活动用户会话', level=1)
    active_users = safe_list(result.get('active_users'))
    if active_users:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = '会话ID'
        hdr[1].text = '登录名'
        hdr[2].text = '主机名'
        hdr[3].text = '程序名'
        hdr[4].text = '状态'
        for item in active_users:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('session_id', ''))
            row_cells[1].text = str(item.get('login_name', ''))
            row_cells[2].text = str(item.get('host_name', ''))
            row_cells[3].text = str(item.get('program_name', ''))
            row_cells[4].text = str(item.get('status', ''))
        set_table_style(table)
    else:
        doc.add_paragraph("无活动用户会话")

    # 七、锁和死锁信息（新增）
    doc.add_heading('七、锁和死锁信息', level=1)
    lock_data = safe_list(result.get('locks'))
    deadlock_data = safe_list(result.get('deadlocks'))
    lock_count = lock_data[0].get('lock_count', 0) if lock_data else 0
    max_wait = lock_data[0].get('max_wait_time', 0) if lock_data else 0
    deadlock_count = deadlock_data[0].get('deadlock_count', 0) if deadlock_data else 0

    lock_table = doc.add_table(rows=3, cols=2)
    lock_table.style = 'Table Grid'
    lock_items = [
        ('当前锁数量', str(lock_count)),
        ('死锁数量 (过去24小时)', str(deadlock_count)),
        ('最长等待时间 (ms)', str(max_wait))
    ]
    for i, (key, val) in enumerate(lock_items):
        lock_table.cell(i, 0).text = key
        lock_table.cell(i, 1).text = str(val)  # 确保字符串
    set_table_style(lock_table)

    # ====== 以下为删除的第八节：数据库磁盘空间信息 ======
    # 已根据要求删除

    # 八、等待统计 (TOP 10)  （原第九节改为第八节）
    doc.add_heading('八、等待统计 (TOP 10)', level=1)
    wait_stats = safe_list(result.get('wait_stats'))
    if wait_stats:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = '等待类型'
        hdr[1].text = '等待时间(ms)'
        hdr[2].text = '等待任务数'
        for item in wait_stats:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('wait_type', ''))
            row_cells[1].text = str(item.get('wait_time_ms', ''))
            row_cells[2].text = str(item.get('waiting_tasks_count', ''))
        set_table_style(table)
    else:
        doc.add_paragraph("无等待统计信息")

    # 九、阻塞进程  （原第十节改为第九节）
    doc.add_heading('九、阻塞进程', level=1)
    blocking = safe_list(result.get('blocking'))
    if blocking:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = '被阻塞会话'
        hdr[1].text = '阻塞会话'
        hdr[2].text = '等待时间(ms)'
        hdr[3].text = '等待类型'
        for item in blocking:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('session_id', ''))
            row_cells[1].text = str(item.get('blocking_session_id', ''))
            row_cells[2].text = str(item.get('wait_duration_ms', ''))
            row_cells[3].text = str(item.get('wait_type', ''))
        set_table_style(table)
    else:
        doc.add_paragraph("无阻塞进程")

    # 十、备份信息  （原第十一节改为第十节）
    doc.add_heading('十、备份信息', level=1)
    backup_info = safe_list(result.get('backup_info'))
    if backup_info:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = '数据库'
        hdr[1].text = '最近备份时间'
        hdr[2].text = '备份类型'
        hdr[3].text = '备份大小(MB)'
        for item in backup_info:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('database_name', ''))
            row_cells[1].text = str(item.get('last_backup_time', ''))
            row_cells[2].text = str(item.get('backup_type', ''))
            row_cells[3].text = str(item.get('backup_size_mb', ''))
        set_table_style(table)
    else:
        doc.add_paragraph("无备份信息（可能权限不足或未备份）")

    # 十一、主机资源使用情况（系统级） （原第十二节改为第十一节）
    doc.add_heading('十一、主机资源使用情况', level=1)
    
    # 磁盘表格
    doc.add_heading('磁盘使用率', level=2)
    host_disk_dict = safe_dict(result.get('host_disk_usage'))
    host_disk = safe_list(host_disk_dict.get('data'))
    if host_disk:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = '挂载点'
        hdr[1].text = '总大小(GB)'
        hdr[2].text = '可用(GB)'
        hdr[3].text = '使用率(%)'
        for row in host_disk:
            if len(row) >= 7:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row[1])
                row_cells[1].text = str(row[3])
                row_cells[2].text = str(row[5])
                row_cells[3].text = str(row[6])
        set_table_style(table)
    else:
        doc.add_paragraph("未获取到磁盘使用率数据")
    
    # 内存表格
    doc.add_heading('内存使用率', level=2)
    host_mem_dict = safe_dict(result.get('host_memory_usage'))
    host_mem = safe_list(host_mem_dict.get('data'))
    if host_mem:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = '内存类型'
        hdr[1].text = '总大小(GB)'
        hdr[2].text = '已用(GB)'
        hdr[3].text = '可用(GB)'
        hdr[4].text = '使用率(%)'
        for mem in host_mem:
            if len(mem) >= 5:
                row_cells = table.add_row().cells
                row_cells[0].text = str(mem[0])
                row_cells[1].text = str(mem[1])
                row_cells[2].text = str(mem[2])
                row_cells[3].text = str(mem[3])
                row_cells[4].text = str(mem[4])
        set_table_style(table)
    else:
        doc.add_paragraph("未获取到内存使用率数据")
    
    # CPU表格
    doc.add_heading('CPU使用率', level=2)
    host_cpu_dict = safe_dict(result.get('host_cpu_usage'))
    host_cpu = safe_list(host_cpu_dict.get('data'))
    if host_cpu:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = '指标'
        hdr[1].text = '值'
        cpu = host_cpu[0]
        row_cells = table.add_row().cells
        row_cells[0].text = 'CPU使用率'
        row_cells[1].text = str(cpu[1])
        set_table_style(table)
    else:
        doc.add_paragraph("未获取到CPU使用率数据")

    # 十二、健康状态汇总  （原第十三节改为第十二节）
    doc.add_heading('十二、健康状态汇总', level=1)
    stats = analyze_inspection_results(result)
    summary = [
        ('整体健康状态', stats['health_status']),
        ('告警总数', stats['alarm_count']),
        ('严重告警', stats['critical_alarm']),
        ('数据库数量', stats['tablespace_count']),
        ('主机内存使用率', f"{stats['host_memory_usage']}%"),
        ('主机CPU使用率', f"{stats['host_cpu_usage']}%"),
        ('磁盘最大使用率', f"{stats['disk_usage_rate']}%"),
        ('状态描述', stats['health_desc'])
    ]
    table = doc.add_table(rows=len(summary), cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(summary):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = str(value)
    set_table_style(table)

    # 添加页脚
    try:
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer_run.font.size = Pt(9)
    except:
        pass

    return doc

# ==================== 生成报告函数（增加详细错误输出） ====================
def generate_report(db_info, result, inspector="", system="", report_type="word"):
    try:
        base_path = os.path.dirname(sys.executable) if getattr(sys, 'frozen', False) else os.path.dirname(os.path.abspath(__file__))
        reports_dir = os.path.join(base_path, 'reports')
        os.makedirs(reports_dir, exist_ok=True)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        label = db_info.get('label', db_info.get('name', 'Unknown')).replace(' ', '_')
        if report_type == "word":
            filename = f"SQLServer_{label}_report_{timestamp}.docx"
            output_path = os.path.join(reports_dir, filename)
            doc = create_sqlserver_report(db_info, result, inspector, system)
            doc.save(output_path)
            printer.print_success(f"Word报告已生成: {output_path}")
            return output_path
        elif report_type == "excel":
            return None
    except Exception as e:
        printer.print_error(f"生成报告失败: {e}")
        traceback.print_exc()
        return None

# ==================== 手动巡检函数 ====================
def single_db_process(sql_templates):
    printer.print_step(1, "手动巡检模式", "start")
    print("\n请输入SQL Server数据库连接信息：")
    label = input("数据库标签 [TEST]: ").strip() or "TEST"
    server = input("服务器地址 [localhost]: ").strip() or "localhost"
    port = input("端口 [1433]: ").strip() or "1433"
    database = input("数据库名称 [master]: ").strip() or "master"
    user = input("用户名 [sa]: ").strip() or "sa"
    import getpass
    password = getpass.getpass("密码: ").strip()

    # 列出可用驱动供选择
    available_drivers = pyodbc.drivers()
    sql_server_drivers = [d for d in available_drivers if 'SQL Server' in d]
    if not sql_server_drivers:
        printer.print_error("未找到任何 SQL Server ODBC 驱动，请先安装。")
        return
    printer.print_info("可用的 SQL Server ODBC 驱动：")
    for i, d in enumerate(sql_server_drivers, 1):
        print(f"  {i}. {d}")
    choice = input(f"请选择驱动编号 [1]: ").strip()
    try:
        idx = int(choice) - 1 if choice else 0
        driver = sql_server_drivers[idx]
    except:
        driver = sql_server_drivers[0]

    # 输入巡检人和巡检系统
    inspector = input("请输入巡检人姓名 [Admin]: ").strip() or "Admin"
    system = input("请输入巡检系统名称 [巡检系统]: ").strip() or "巡检系统"

    db_info = {
        'label': label,
        'server': server,
        'port': port,
        'name': database,
        'user': user,
        'password': password,
        'driver': driver
    }

    printer.print_info("正在连接数据库...")
    try:
        data_getter = getData(server, port, user, password, driver, label)
    except Exception as e:
        printer.print_error(f"连接失败: {e}")
        return

    printer.print_info("执行巡检查询...")
    result = data_getter.checkdb(sql_templates)

    result['host_disk_usage'] = get_host_disk_usage()
    result['host_memory_usage'] = get_host_memory_usage()
    result['host_cpu_usage'] = get_host_cpu_usage()

    stats = analyze_inspection_results(result)
    printer.print_success(f"健康状态: {stats['health_status']} - {stats['health_desc']}")

    report_path = generate_report(db_info, result, inspector, system, report_type="word")
    if report_path:
        printer.print_success(f"报告生成成功: {report_path}")

# ==================== 批量巡检函数 ====================
def batch_process(sql_templates, excel_file, summary_report):
    printer.print_step(1, "批量巡检模式", "start")
    connections = read_connections_from_excel(excel_file)
    if not connections:
        printer.print_error("未读取到任何有效连接，退出")
        return

    # 批量模式下，巡检人和系统可统一输入
    inspector = input("请输入巡检人姓名 [Admin]: ").strip() or "Admin"
    system = input("请输入巡检系统名称 [巡检系统]: ").strip() or "巡检系统"

    summary_data = []
    for idx, conn in enumerate(connections, 1):
        printer.print_step(2, f"处理第 {idx}/{len(connections)} 个数据库: {conn['Label']}", "start")
        available_drivers = pyodbc.drivers()
        sql_server_drivers = [d for d in available_drivers if 'SQL Server' in d]
        if not sql_server_drivers:
            printer.print_error("未找到任何 SQL Server ODBC 驱动，跳过该数据库")
            continue
        driver = conn['Driver'].strip('{}')
        if driver not in sql_server_drivers:
            matched = [d for d in sql_server_drivers if driver.lower() in d.lower()]
            if matched:
                printer.print_warning(f"驱动 '{conn['Driver']}' 不匹配，使用 '{matched[0]}'")
                driver = matched[0]
            else:
                printer.print_error(f"驱动 '{conn['Driver']}' 无效，且无匹配项，跳过")
                continue

        db_info = {
            'label': conn['Label'],
            'server': conn['Server'],
            'port': conn['Port'],
            'name': conn['Database'],
            'user': conn['User'],
            'password': conn['Password'],
            'driver': driver
        }
        try:
            data_getter = getData(db_info['server'], db_info['port'], db_info['user'],
                                  db_info['password'], db_info['driver'], db_info['label'])
        except Exception as e:
            printer.print_error(f"连接失败: {e}")
            continue

        result = data_getter.checkdb(sql_templates)
        result['host_disk_usage'] = get_host_disk_usage()
        result['host_memory_usage'] = get_host_memory_usage()
        result['host_cpu_usage'] = get_host_cpu_usage()

        report_path = generate_report(db_info, result, inspector, system, report_type="word")
        if report_path and summary_report:
            stats = analyze_inspection_results(result)
            summary_item = {
                'label': db_info['label'],
                'server': db_info['server'],
                'database': db_info['name'],
                'version': result.get('version', [{}])[0].get('version', '')[:50] if result.get('version') else '',
                'total_size_mb': sum([item.get('size_mb', 0) for item in result.get('db_size', [])]),
                'active_sessions': len(result.get('active_users', [])),
                'max_connections': result.get('max_connections', [{}])[0].get('max_connections', '') if result.get('max_connections') else '',
                'check_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'report_file': report_path
            }
            summary_data.append(summary_item)

    if summary_report and summary_data:
        base_path = os.path.dirname(sys.executable) if getattr(sys, 'frozen', False) else os.path.dirname(os.path.abspath(__file__))
        reports_dir = os.path.join(base_path, 'reports')
        os.makedirs(reports_dir, exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        summary_file = os.path.join(reports_dir, f"summary_report_{timestamp}.xlsx")
        generate_summary_excel(summary_data, summary_file)

# ==================== 主菜单 ====================
def show_menu():
    printer.print_step(0, "SQL Server 数据库巡检工具", "start")
    print("1、手动巡检")
    print("2、批量巡检")
    print("3、生成批量巡检模板")
    print("4、结束巡检")
    choice = input("请输入选择 (1-4): ").strip()
    return choice

def main():
    printer.print_step(0, "许可证验证", "start")
    try:
        validator = LicenseValidator()
        is_valid, message, remaining_days = validator.validate_license()
        if is_valid:
            printer.print_success(message)
            if remaining_days < 30:
                printer.print_warning(f"试用期即将结束，剩余 {remaining_days} 天")
        else:
            printer.print_error(message)
            if "试用期" in message and remaining_days <= 0:
                printer.print_error("试用期已结束，请购买正式版许可证")
                sys.exit(1)
            elif "许可证文件不存在" in message:
                printer.print_info("已自动创建试用许可证")
            else:
                printer.print_error("许可证验证失败，程序退出")
                sys.exit(1)
    except Exception as e:
        printer.print_error(f"许可证验证失败: {e}")

    args = passArgu().get_argus()

    if args.generate_excel_template:
        generate_excel_template(args.generate_excel_template)
        sys.exit(0)

    sql_templates = get_resource_path(args.sqltemplates)
    # 始终覆盖初始化模板，确保最新版本
    init_sql_templates(sql_templates)

    if args.excel_connections:
        batch_process(sql_templates, args.excel_connections, args.summary_report)
        return

    while True:
        choice = show_menu()
        if choice == '1':
            single_db_process(sql_templates)
        elif choice == '2':
            excel_file = input("请输入Excel连接文件路径: ").strip()
            if not excel_file:
                printer.print_error("文件路径不能为空")
                continue
            if not os.path.exists(excel_file):
                printer.print_error("文件不存在")
                continue
            summary_choice = input("是否生成汇总Excel报告? (y/n) [n]: ").strip().lower()
            summary_report = (summary_choice == 'y')
            batch_process(sql_templates, excel_file, summary_report)
        elif choice == '3':
            template_file = input("请输入要保存的Excel模板路径 [connections_template.xlsx]: ").strip()
            if not template_file:
                template_file = "connections_template.xlsx"
            generate_excel_template(template_file)
        elif choice == '4':
            printer.print_success("再见！")
            break
        else:
            printer.print_error("无效选择，请重新输入")

        input("\n按回车键继续...")

if __name__ == "__main__":
    main()